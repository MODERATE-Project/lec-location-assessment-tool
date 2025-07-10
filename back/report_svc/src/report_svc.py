from flask import Flask, jsonify, request, make_response, send_file
from flask_cors import CORS
import docx
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import requests
import os

import fields.field_manager as field_manager
import logging

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "DEBUG"),
    # Formato del mensaje
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',  # Formato de la fecha
    handlers=[
        logging.StreamHandler()  # Enviar los logs a la consola (stdout)
    ]
)

log = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

BASE_DIR = '/report_svc/data'
REPORT_FILE_FULL = os.getenv("REPORT_FILE", 'report_template.docx')
REPORT_FILE_AREA_SELECTED = os.getenv(
    "REPORT_FILE_AREA_SELECTED", 'report_template_area_selected.docx')


# buildings = get_buildings_data()

# @app.route('/buildings/old', methods=['GET'])
# def get_buildings():
#     municipio = request.args.get('municipio', None)  # Obtiene el parámetro municipio; si no está presente, retorna None
#     # coords = request.json.get('coordinates', [])

#     filtered_buildings = [building for building in buildings if building['Municipios'] == municipio]
#     return jsonify({"buildings": filtered_buildings})

# def create_dynamic_table(paragraph, table_data):
#     """Helper function to create a dynamic table based on provided data"""
#     rows_needed = len(table_data["rows"]) + 1          # +1 for header row
#     cols_needed = len(table_data["headers"])

#     # 1. insert table *before* the placeholder paragraph
#     table = paragraph._parent.add_table(rows=rows_needed, cols=cols_needed, width=Inches(6))

#     # 2. populate header row
#     for col_idx, header in enumerate(table_data["headers"]):
#         table.rows[0].cells[col_idx].text = str(header)

#     # 3. populate the data rows
#     for row_idx, row_data in enumerate(table_data["rows"], start=1):
#         for col_idx, value in enumerate(row_data):
#             table.rows[row_idx].cells[col_idx].text = str(value)

#     # 4. remove the placeholder paragraph
#     # paragraph._element.getparent().remove(paragraph._element) 


def build_table_from_dict(table_dict):
    """Crea una tabla de docx a partir de un diccionario con 'headers' y 'rows'."""
    temp_doc = docx.Document()  # documento temporal
    section = temp_doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin

    headers = table_dict.get("headers", [])
    rows = table_dict.get("rows", [])

    cols = len(headers)
    table = temp_doc.add_table(rows=1, cols=cols)
    table.autofit = False
    col_width = page_width / cols


    # Encabezados
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = str(header)
        cell.width = col_width
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(10)

        set_cell_background_color(cell, "FFFFFF")  # Fondo blanco
        set_cell_border(cell)

    # Filas alternas
    for row_idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        bg_color = "f2f2f2" if row_idx % 2 == 0 else "FFFFFF"

        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.text = str(val)
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.size = Pt(10)

            set_cell_background_color(cell, bg_color)
            set_cell_border(cell)

    return table

def set_cell_border(cell):
    """Agrega bordes visibles a la celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)

    for border_name in ["top", "left", "bottom", "right"]:
        element = OxmlElement(f"w:{border_name}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "c4c4c4")
        borders.append(element)

def set_cell_background_color(cell, color):
    """Establece el color de fondo de una celda (hex sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def insert_table_after_paragraph(paragraph, table_to_insert):
    # new_table = paragraph._parent.add_table(rows=0, cols=0)
    # new_tbl = new_table._tbl
    # tbl_to_copy = table_to_insert._tbl
    # new_tbl.clear_content()
    # new_tbl.extend(tbl_to_copy)

    # body = paragraph._parent
    # body.remove(new_tbl)
    # ref_element = paragraph._element
    # body.insert(index + 1, new_tbl)

    tbl_to_copy = table_to_insert._tbl
    new_tbl = tbl_to_copy.__copy__()
    paragraph._element.addnext(new_tbl)



def replace_text_in_paragraphs(paragraphs, data):
    """Helper function to replace placeholders in regular paragraphs"""

    for i, paragraph in enumerate(paragraphs):
        for key, value in data.items():
            key_pattern = f"${{{key}}}"
            if key_pattern in paragraph.text:

                if isinstance(value, dict) and 'table' in value:
                    table_data = value['table']

                    # Si es un dict con headers/rows, lo convertimos en una tabla real
                    if isinstance(table_data, dict):
                        table_to_insert = build_table_from_dict(table_data)
                    else:
                        table_to_insert = table_data


                    # Insertar la tabla justo después
                    insert_table_after_paragraph(paragraph, table_to_insert)

                    # Eliminar el párrafo con el marcador
                    p = paragraph._element
                    p.getparent().remove(p)

                elif "IMG" in key or "PLOT" in key:

                    # Clear the placeholder text
                    paragraph.text = ""

                    # Add the plot image to this paragraph
                    run = paragraph.add_run()

                    log.info(f"embbeding img_file: {value}")
                    try:
                        with open(os.path.join(BASE_DIR, value), 'rb') as image_stream:
                            run.add_picture(image_stream, width=Inches(5.0))
                    except (FileNotFoundError, TypeError):
                        print(f"Error: No se pudo encontrar el archivo de imagen {value}")

                else:
                    if app.debug == True:
                        paragraph.text = paragraph.text.replace(key, str(value))  # FIXME: This is used for testing
                    else:
                        paragraph.text = paragraph.text.replace(key_pattern, str(value)) # FIXME: This is the correct final way


# def replace_text_in_tables(tables, data):
#     """Helper function to replace placeholders in tables"""
#     for table in tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     for key, value in data.items():
#                         key_pattern = f"${{{key}}}"
#                         if key_pattern in paragraph.text:
#                             paragraph.text = paragraph.text.replace(key_pattern, str(value))


def replace_text_in_headers_footers(sections, data):
    """Helper function to replace placeholders in headers and footers"""
    for section in sections:
        header = section.header
        footer = section.footer

        replace_text_in_paragraphs(header.paragraphs, data)
        replace_text_in_paragraphs(footer.paragraphs, data)


@app.route('/report', methods=['POST'])
def get_report():

    data = request.json

    # data = { '${MUNICIPALITY_TITLE}':'CREVILLENT',
    #         '${MUNICIPALITY}': 'Crevillent',
    #         '${NUM_BUILDINGS}': '12',
    #         '${PCT_1} ': '69',
    #         '${PCT_4} ': '13',
    #         '${PCT_5} ': '10',
    #         '${PCT_6} ': '8' 
    #         }

    print('aqui van los datos', data)

    municipality = data.get("MUNICIPALITY").capitalize()

    municipality_parameters = field_manager.get_and_compute_as_needed(
        municipality=municipality, field_dict=data, base_dir=BASE_DIR)

    isAreaSelected = bool(data.get("isAreaSelected"))
    if isAreaSelected:
        report_file = REPORT_FILE_AREA_SELECTED
    else:
        report_file = REPORT_FILE_FULL
    log.debug(f"Using report template file {report_file} because isAreaSelected is {isAreaSelected}")

    doc_path = os.path.join(BASE_DIR, report_file)
    report_filled_path = os.path.join(BASE_DIR, 'generated_report.docx')
    pdf_path = os.path.join(BASE_DIR, 'generated_report.pdf')

    # convert_file(doc_path, 'pdf', outputfile=pdf_path)
    # convert_odt_to_pdf(doc_path, pdf_path)

    # Load DOCX template
    doc = docx.Document(doc_path)

    replace_text_in_paragraphs(doc.paragraphs, municipality_parameters)

    # Replace text in tables
    # replace_text_in_tables(doc.tables, municipality_parameters)

    # Replace placeholders in headers and footers
    replace_text_in_headers_footers(doc.sections, municipality_parameters)

    doc.save(report_filled_path)

    # document = Document()
#
    # Load a doc or docx file
    # document.LoadFromFile(report_filled_path)

    # Save the document to PDF
    # document.SaveToFile(pdf_path, FileFormat.PDF)
    # document.Close()

    # Abre el archivo en modo binario
    with open(report_filled_path, 'rb') as file:
        files = {'document': file}

        # Realiza la solicitud POST al contenedor 'docx-to-pdf'
        response = requests.post("http://docx-to-pdf:8080/pdf", files=files)

        if response.ok:

            with open(pdf_path, 'wb') as result_file:
                result_file.write(response.content)

            return send_file(pdf_path, as_attachment=True, download_name=f"report_{data['MUNICIPALITY']}.pdf", mimetype='application/pdf')

        return f"Error in processing: {response.status_code}", 500

    # return send_file(report_filled_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
